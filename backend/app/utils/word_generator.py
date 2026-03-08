import os
import requests
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_border(cell, **kwargs):
    """
    Set cell border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "single"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, create
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, create
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def style_border_none(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
        
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def render_word_from_invoice(invoice, services, total, provider, customization=None):
    doc = Document()
    
    # Page setup
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    HEADER_GREEN = RGBColor(0x1B, 0x43, 0x32) # #1b4332
    LIGHT_GREEN_BG = 'F4F9F7'

    # --- 1. HEADER (Logo and Title) ---
    header_table = doc.add_table(rows=1, cols=2)
    header_table.width = Inches(7.5)
    style_border_none(header_table)
    
    # Logo
    if customization and customization.logo_url and customization.include_logo:
        try:
            response = requests.get(customization.logo_url, timeout=5)
            if response.status_code == 200:
                img_stream = BytesIO(response.content)
                cell_logo = header_table.cell(0, 0)
                paragraph = cell_logo.paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(img_stream, width=Inches(1.6))
        except:
            pass
            
    # Title and Number
    cell_title = header_table.cell(0, 1)
    title_pg = cell_title.paragraphs[0]
    title_pg.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    run_h1 = title_pg.add_run(f"CUENTA DE COBRO N° {invoice.invoice_number}")
    run_h1.font.name = 'Open Sans'
    run_h1.font.size = Pt(18)
    run_h1.font.bold = True
    run_h1.font.color.rgb = HEADER_GREEN
    
    date_pg = cell_title.add_paragraph()
    date_pg.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    invoice_date = invoice.invoice_date.strftime('%d/%m/%Y') if hasattr(invoice.invoice_date, 'strftime') else str(invoice.invoice_date)
    run_date = date_pg.add_run(f"Fecha: {invoice_date}")
    run_date.font.size = Pt(11)
    run_date.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    doc.add_paragraph()

    # --- 2. INFO BOX ---
    info_table = doc.add_table(rows=3, cols=2)
    style_border_none(info_table)
    # Set column widths
    info_table.columns[0].width = Inches(1.5)
    info_table.columns[1].width = Inches(5.5)
    
    def style_info_row(row_idx, label, value):
        cell_label = info_table.cell(row_idx, 0)
        cell_label.text = label
        paragraph = cell_label.paragraphs[0]
        paragraph.paragraph_format.space_before = Pt(4)
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.runs[0]
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = HEADER_GREEN
        
        cell_value = info_table.cell(row_idx, 1)
        cell_value.text = value
        val_pg = cell_value.paragraphs[0]
        val_pg.paragraph_format.space_before = Pt(4)
        val_pg.paragraph_format.space_after = Pt(4)
        val_pg.runs[0].font.size = Pt(10)
        
    style_info_row(0, "CLIENTE:", invoice.service_account.client.name)
    
    start_date = invoice.service_account.start_date.strftime('%d/%m/%Y') if hasattr(invoice.service_account.start_date, 'strftime') else str(invoice.service_account.start_date)
    end_date = invoice.service_account.end_date.strftime('%d/%m/%Y') if hasattr(invoice.service_account.end_date, 'strftime') else str(invoice.service_account.end_date)
    style_info_row(1, "PERIODO:", f"{start_date} — {end_date}")
    
    service_txt = customization.service_text if customization and customization.service_text else "SERVICIO DE TRANSPORTE"
    style_info_row(2, "SERVICIO:", service_txt)

    doc.add_paragraph()

    # --- 3. SERVICES TABLE ---
    services_table = doc.add_table(rows=1, cols=5)
    services_table.style = 'Table Grid'
    
    # Headers
    h_cells = services_table.rows[0].cells
    headers = ["FECHA", "SERVICIO / MATERIAL", "CANT", "V. UNITARIO", "TOTAL"]
    for i, h in enumerate(headers):
        cell = h_cells[i]
        cell.text = h
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '1B4332') # Use hex string directly
        cell._tc.get_or_add_tcPr().append(shading_elm)
        
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.runs[0]
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data
    for i, s in enumerate(services):
        row_cells = services_table.add_row().cells
        s_date = s.service_date.strftime('%d/%m/%Y') if hasattr(s.service_date, 'strftime') else str(s.service_date)
        
        if s.custom_material:
            s_name = f"Viaje de {s.custom_material}"
        elif s.material:
            s_name = f"Viaje de {s.material.name}"
        else:
            s_name = "Viaje sin material"
            
        row_cells[0].text = s_date
        row_cells[1].text = s_name
        row_cells[2].text = str(s.quantity)
        row_cells[3].text = f"${s.price:,.0f}" if s.price else "$0"
        row_cells[4].text = f"${s.total_amount:,.0f}" if s.total_amount else "$0"
        
        # Zebra striping
        if i % 2 == 1:
            for cell in row_cells:
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'F9FAF9')
                cell._tc.get_or_add_tcPr().append(shading_elm)

        for j, cell in enumerate(row_cells):
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if j != 1 else WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.runs[0]
            run.font.size = Pt(9)

    doc.add_paragraph()

    # --- 4. PAYMENT AND SIGNATURE ---
    footer_table = doc.add_table(rows=1, cols=2)
    footer_table.width = Inches(7.5)
    style_border_none(footer_table)
    
    # Left: Bank Info & Total
    cell_left = footer_table.cell(0, 0)
    
    total_pg = cell_left.add_paragraph()
    run_total_lbl = total_pg.add_run("TOTAL A PAGAR: ")
    run_total_lbl.font.bold = True
    run_total_lbl.font.size = Pt(12)
    run_total_lbl.font.color.rgb = HEADER_GREEN
    
    run_total_val = total_pg.add_run(f"${total:,.0f}")
    run_total_val.font.bold = True
    run_total_val.font.size = Pt(16)
    run_total_val.font.color.rgb = HEADER_GREEN

    if customization and customization.include_bank_info and customization.provider_bank:
        bank_pg = cell_left.add_paragraph()
        bank_pg.paragraph_format.space_before = Pt(10)
        run_bank = bank_pg.add_run(f"Banco: {customization.provider_bank}\n"
                                   f"Tipo: {customization.provider_type_account}\n"
                                   f"Cuenta: {customization.provider_number_account}")
        run_bank.font.size = Pt(10)
        run_bank.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Right: Signature
    cell_right = footer_table.cell(0, 1)
    cell_right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    if customization and customization.signature_url and customization.include_signature:
        try:
            response = requests.get(customization.signature_url, timeout=5)
            if response.status_code == 200:
                sig_stream = BytesIO(response.content)
                paragraph = cell_right.paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(sig_stream, width=Inches(1.6))
        except:
            pass
            
    name_pg = cell_right.add_paragraph()
    name_pg.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_name = name_pg.add_run(provider.name)
    run_name.font.size = Pt(14)
    run_name.font.bold = True
    run_name.font.italic = True
    run_name.font.color.rgb = HEADER_GREEN
    
    if provider.document_number:
        doc_pg = cell_right.add_paragraph()
        doc_pg.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_doc = doc_pg.add_run(f"CC: {provider.document_number}")
        run_doc.font.size = Pt(11)
        run_doc.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Footer message
    if customization and customization.include_footer and customization.footer_message:
        doc.add_paragraph()
        footer_pg = doc.add_paragraph(customization.footer_message)
        footer_pg.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_footer = footer_pg.runs[0]
        run_footer.font.italic = True
        run_footer.font.size = Pt(11)
        run_footer.font.color.rgb = HEADER_GREEN

    # Save
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
